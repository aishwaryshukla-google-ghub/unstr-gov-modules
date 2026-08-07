import io
from .base import BaseDataHandler


class DocxHandler(BaseDataHandler):
    """
    Handler for Microsoft Word (.docx) documents.
    Extracts paragraphs, headings, and tables and renders them into structured PDF pages.
    """

    def convert_to_pdf(self, file_bytes: bytes, filename: str) -> bytes:
        try:
            import docx
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

            doc_in = docx.Document(io.BytesIO(file_bytes))

            buffer = io.BytesIO()
            doc_out = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=40,
                leftMargin=40,
                topMargin=40,
                bottomMargin=40
            )

            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "DocxTitle",
                parent=styles["Heading1"],
                fontSize=16,
                leading=20,
                textColor=colors.HexColor("#1a73e8"),
                spaceAfter=12
            )
            h2_style = ParagraphStyle(
                "DocxH2",
                parent=styles["Heading2"],
                fontSize=13,
                leading=16,
                textColor=colors.HexColor("#202124"),
                spaceBefore=8,
                spaceAfter=4
            )
            body_style = ParagraphStyle(
                "DocxBody",
                parent=styles["Normal"],
                fontSize=10,
                leading=14,
                textColor=colors.HexColor("#202124"),
                spaceAfter=6
            )
            cell_style = ParagraphStyle(
                "DocxCell",
                parent=styles["Normal"],
                fontSize=8,
                leading=10,
                textColor=colors.HexColor("#202124")
            )

            story = [Paragraph(f"Document: {filename}", title_style), Spacer(1, 8)]

            # 1. Process paragraphs
            for p in doc_in.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if p.style.name.startswith("Heading"):
                    story.append(Paragraph(safe_text, h2_style))
                else:
                    story.append(Paragraph(safe_text, body_style))

            # 2. Process tables
            for table in doc_in.tables:
                story.append(Spacer(1, 6))
                t_data = []
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        c_text = cell.text.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        row_cells.append(Paragraph(c_text, cell_style))
                    t_data.append(row_cells)

                if t_data:
                    t = Table(t_data)
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f3f4")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dadce0")),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 8))

            doc_out.build(story)
            return buffer.getvalue()

        except ImportError:
            # Fallback simple text-based rendering
            from .text_md_handler import TextMarkdownHandler
            return TextMarkdownHandler().convert_to_pdf(b"Word document content (python-docx not installed).", filename)
